import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import cm
import os
from datetime import datetime

class ExportManager:
    def __init__(self, db_manager):
        self.db_manager = db_manager
    
    def safe_execute_query(self, query, params=None):
        """تنفيذ استعلام بشكل آمن مع معالجة الأخطاء"""
        try:
            result = self.db_manager.execute_query(query, params)
            return result if result is not None else []
        except Exception as e:
            print(f"خطأ في تنفيذ الاستعلام: {e}")
            return []
    
    def arabic_text(self, text):
        """معالجة النص العربي للعرض الصحيح"""
        if text is None:
            return ""
        return str(text)
    
    def safe_get_column(self, row, index, default=""):
        """الحصول على عمود من الصف بشكل آمن"""
        if row is None or index >= len(row):
            return default
        return row[index] if row[index] is not None else default
    
    def export_to_excel(self, data, columns, file_path, sheet_name="بيانات"):
        """تصدير البيانات إلى Excel"""
        try:
            # التأكد من أن البيانات ليست None
            if data is None:
                data = []
            
            # معالجة البيانات
            processed_data = []
            for row in data:
                if row is not None:
                    processed_row = [str(cell) if cell is not None else "" for cell in row]
                    processed_data.append(processed_row)
            
            # التأكد من وجود بيانات
            if not processed_data:
                processed_data = [["لا توجد بيانات" for _ in columns]]
            
            df = pd.DataFrame(processed_data, columns=columns)
            
            # إنشاء كاتب Excel
            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name=sheet_name, index=False)
                
                # الحصول على ورقة العمل وتعديلها
                worksheet = writer.sheets[sheet_name]
                
                # ضبط اتجاه النص للخلايا
                for row in worksheet.iter_rows():
                    for cell in row:
                        cell.alignment = cell.alignment.copy(
                            horizontal='right',
                            vertical='center'
                        )
                
                # ضبط عرض الأعمدة تلقائياً
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
            
            return True
        except Exception as e:
            print(f"خطأ في التصدير إلى Excel: {e}")
            return False
    
    def export_to_word(self, data, columns, file_path, title="تقرير"):
        """تصدير البيانات إلى Word"""
        try:
            # التأكد من أن البيانات ليست None
            if data is None:
                data = []
            
            doc = Document()
            
            # العنوان
            doc.add_heading(title, level=1)
            
            # تاريخ التصدير
            date_text = f"تاريخ التصدير: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            doc.add_paragraph(date_text)
            doc.add_paragraph()
            
            # التحقق من وجود بيانات
            if not data:
                doc.add_paragraph("لا توجد بيانات لعرضها")
                doc.save(file_path)
                return True
            
            # إنشاء الجدول
            table = doc.add_table(rows=1, cols=len(columns))
            table.style = 'Light Grid'
            
            # رأس الجدول
            header_cells = table.rows[0].cells
            for i, column in enumerate(columns):
                header_cells[i].text = str(column)
            
            # بيانات الجدول
            for row in data:
                if row is not None:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row):
                        row_cells[i].text = str(cell) if cell is not None else ""
            
            doc.save(file_path)
            return True
        except Exception as e:
            print(f"خطأ في التصدير إلى Word: {e}")
            return False
    
    def export_to_pdf(self, data, columns, file_path, title="تقرير"):
        """تصدير البيانات إلى PDF"""
        try:
            # التأكد من أن البيانات ليست None
            if data is None:
                data = []
            
            doc = SimpleDocTemplate(
                file_path, 
                pagesize=A4,
                rightMargin=2*cm,
                leftMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )
            elements = []
            styles = getSampleStyleSheet()
            
            # العنوان
            title_paragraph = Paragraph(title, styles['Heading1'])
            elements.append(title_paragraph)
            elements.append(Spacer(1, 12))
            
            # تاريخ التصدير
            date_text = f"تاريخ التصدير: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            date_paragraph = Paragraph(date_text, styles['Normal'])
            elements.append(date_paragraph)
            elements.append(Spacer(1, 20))
            
            # التحقق من وجود بيانات
            if not data:
                no_data_text = "لا توجد بيانات لعرضها"
                no_data_paragraph = Paragraph(no_data_text, styles['Normal'])
                elements.append(no_data_paragraph)
                doc.build(elements)
                return True
            
            # تحضير بيانات الجدول
            table_data = []
            
            # رأس الجدول
            header_row = [str(col) for col in columns]
            table_data.append(header_row)
            
            # بيانات الجدول
            for row in data:
                if row is not None:
                    processed_row = [str(cell) if cell is not None else "" for cell in row]
                    table_data.append(processed_row)
            
            # إنشاء الجدول
            table = Table(table_data, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E86AB')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8F9FA')),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#DEE2E6')),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            
            elements.append(table)
            doc.build(elements)
            return True
        except Exception as e:
            print(f"خطأ في التصدير إلى PDF: {e}")
            return False
    
    def generate_incoming_report(self, start_date=None, end_date=None):
        """إنشاء تقرير الوارد"""
        try:
            query = """
            SELECT 
                ir.record_number, 
                ir.incoming_number, 
                ir.serial_number, 
                ir.title,
                isrc.name as source, 
                it.name as type, 
                e.name as employee,
                s.name as specialization, 
                ir.registration_date
            FROM incoming_records ir
            LEFT JOIN incoming_sources isrc ON ir.incoming_source_id = isrc.id
            LEFT JOIN incoming_types it ON ir.incoming_type_id = it.id
            LEFT JOIN employees e ON ir.employee_id = e.id
            LEFT JOIN specializations s ON ir.specialization_id = s.id
            """
            
            params = []
            if start_date and end_date:
                query += " WHERE ir.registration_date BETWEEN ? AND ?"
                params.extend([start_date, end_date])
            
            query += " ORDER BY ir.registration_date DESC"
            
            data = self.safe_execute_query(query, params)
            columns = [
                'رقم السجل', 'رقم الوارد', 'الرقم التسلسلي', 'العنوان', 
                'جهة الوارد', 'النوع', 'الموظف', 'الاختصاص', 'تاريخ التسجيل'
            ]
            
            return data, columns
        except Exception as e:
            print(f"خطأ في إنشاء تقرير الوارد: {e}")
            return [], []
    
    def generate_outgoing_report(self, start_date=None, end_date=None):
        """إنشاء تقرير الصادر"""
        try:
            query = """
            SELECT 
                orc.record_number, 
                orc.outgoing_number, 
                orc.serial_number, 
                orc.title,
                od.name as destination, 
                e.name as employee,
                s.name as specialization, 
                orc.registration_date
            FROM outgoing_records orc
            LEFT JOIN outgoing_destinations od ON orc.outgoing_destination_id = od.id
            LEFT JOIN employees e ON orc.employee_id = e.id
            LEFT JOIN specializations s ON orc.specialization_id = s.id
            """
            
            params = []
            if start_date and end_date:
                query += " WHERE orc.registration_date BETWEEN ? AND ?"
                params.extend([start_date, end_date])
            
            query += " ORDER BY orc.registration_date DESC"
            
            data = self.safe_execute_query(query, params)
            columns = [
                'رقم السجل', 'رقم الصادر', 'الرقم التسلسلي', 'العنوان', 
                'جهة الصادر', 'الموظف', 'الاختصاص', 'تاريخ التسجيل'
            ]
            
            return data, columns
        except Exception as e:
            print(f"خطأ في إنشاء تقرير الصادر: {e}")
            return [], []
    
    def generate_employee_report(self, start_date=None, end_date=None, department=None, status_filter='نشط فقط'):
        """إنشاء تقرير الموظفين"""
        try:
            # بناء استعلام الموظفين
            employee_query = """
            SELECT 
                id, name, department, position, is_active
            FROM employees 
            WHERE 1=1
            """
            params = []
            
            if department and department != 'الكل':
                employee_query += " AND department = ?"
                params.append(department)
            
            if status_filter == 'نشط فقط':
                employee_query += " AND is_active = 1"
            elif status_filter == 'غير نشط فقط':
                employee_query += " AND is_active = 0"
            
            employee_query += " ORDER BY name"
            
            employees = self.safe_execute_query(employee_query, params)
            
            # حساب الإحصائيات التفصيلية
            processed_data = []
            total_faxes = 0
            total_emails = 0
            
            for emp in employees:
                if emp is None:
                    continue
                    
                emp_id = self.safe_get_column(emp, 0, 0)
                
                # حساب الفاكسات (سجلات الصادر)
                fax_query = """
                SELECT COUNT(*) FROM outgoing_records 
                WHERE employee_id = ?
                """
                fax_params = [emp_id]
                
                # حساب الإيميلات (سجلات الوارد من نوع إيميل)
                email_query = """
                SELECT COUNT(*) FROM incoming_records 
                WHERE employee_id = ? 
                AND incoming_type_id IN (SELECT id FROM incoming_types WHERE name LIKE '%إيميل%' OR name LIKE '%email%')
                """
                email_params = [emp_id]
                
                if start_date and end_date:
                    fax_query += " AND registration_date BETWEEN ? AND ?"
                    email_query += " AND registration_date BETWEEN ? AND ?"
                    fax_params.extend([start_date, end_date])
                    email_params.extend([start_date, end_date])
                
                fax_count_result = self.safe_execute_query(fax_query, fax_params)
                email_count_result = self.safe_execute_query(email_query, email_params)
                
                fax_count = fax_count_result[0][0] if fax_count_result and fax_count_result[0] else 0
                email_count = email_count_result[0][0] if email_count_result and email_count_result[0] else 0
                
                total_comm = fax_count + email_count
                
                total_faxes += fax_count
                total_emails += email_count
                
                # حساب النسبة المئوية
                total_all = total_faxes + total_emails
                percentage = (total_comm / total_all * 100) if total_all > 0 else 0
                
                status_text = "نشط" if self.safe_get_column(emp, 4, 0) else "غير نشط"
                
                processed_data.append([
                    self.safe_get_column(emp, 1, "غير معروف"),  # الاسم
                    self.safe_get_column(emp, 2, "غير محدد"),  # القسم
                    self.safe_get_column(emp, 3, "غير محدد"),  # المنصب
                    fax_count,
                    email_count,
                    total_comm,
                    f"{percentage:.1f}%",
                    status_text
                ])
            
            columns = [
                'اسم الموظف', 'القسم', 'المنصب',
                'عدد الفاكسات', 'عدد الإيميلات', 'المجموع', 'النسبة %', 'الحالة'
            ]
            
            return processed_data, columns, total_faxes, total_emails
            
        except Exception as e:
            print(f"خطأ في إنشاء تقرير الموظفين: {e}")
            return [], [], 0, 0